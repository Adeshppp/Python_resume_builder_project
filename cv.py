from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
      pyttsx3.speak(text)

document= Document()

# profile picture
document.add_picture('picture.jpg',width=Inches(2.0))

# name phone number and email details
name = input('What is your NAME? ')
speak('Hello'+ name + 'How are you today?')
speak('What is your phone number?')
phone = input('What is your PHONE? ')
speak('Enter Your Email Address')
email= input('What is your EMAIL? ')

document.add_paragraph(name + ' | ' + phone + ' | '+ email)

# about me 

document.add_heading('About Me')
about_me = input('Tell me about yourself? ')
speak('Tell me about yourself')
document.add_paragraph(about_me)


# work Experiences

document.add_heading('Work Experience')
speak('do you have any Work Experience, if yes then enter otherwise press enter')

p = document.add_paragraph()

company = input('Enter name of the Company')
from_date= input('From date :')
to_date=input('to_date : ')
p.add_run(company + ' ').bold =True
p.add_run(from_date + '-'+ to_date+'\n').italic= True
speak('Describe your experience at '+ company)

experience_details=input('Describe your experience at '+ company)
p.add_run(experience_details)

while True:
      has_more_experiences =input('Do you have more Experiences? Yes or No '+' ')
      if has_more_experiences.lower() == 'yes':
            p = document.add_paragraph()
            company = input('Enter name of the Company')
            from_date= input('From date :')
            to_date=input('to_date : ')
            p.add_run(company + ' ').bold =True
            p.add_run(from_date + '-'+ to_date+'\n').italic= True
            experience_details=input('Describe your experience at '+ company+' ')
            p.add_run(experience_details)
      else:
            break

# SKILLS

document.add_heading('Skills')

while True:
      more_skills=input('Do you want to add skills? Yes or No ')
      if more_skills.lower()=='yes':
            s=document.add_paragraph()
            skills= input('enter your skill : ')
            s.add_run(skills)
            s.style='List Bullet'
      else:
            break

document.save('cv.docx')

